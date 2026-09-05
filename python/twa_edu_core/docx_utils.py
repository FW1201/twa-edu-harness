"""Word 文件的共用版面元件。

原本這份程式碼以完全相同的內容散落在 15 個 skill 的 `scripts/` 底下。
改中文字型或表格底色要改 15 次，漏改一次就產出外觀不一致的 .docx。
"""
from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .fonts import DEFAULT_CJK_FONT, set_east_asia_font
from .theme import (
    BLUE_DEEP,
    BLUE_LIGHT,
    BLUE_MID,
    DARK_TEXT,
    GRAY_LIGHT,
    MUTED_TEXT,
    WHITE,
    rgb_hex,
)


def set_cell_bg(cell, color: RGBColor) -> None:
    """設定儲存格底色。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex(color))
    tcPr.append(shd)


def set_cell_border(cell, color: str = "2471A3", size: str = "4") -> None:
    """為儲存格四邊加上單線框。"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def cell_write(cell, text, bold: bool = False, size: int = 11,
               color: RGBColor = DARK_TEXT, center: bool = False,
               font: str = DEFAULT_CJK_FONT):
    """清空儲存格並寫入一段文字。"""
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = font
    r.font.color.rgb = color
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_east_asia_font(r, font)
    return p


def header_cell(cell, text, bg: RGBColor = BLUE_MID, size: int = 11) -> None:
    """表頭儲存格：深色底、白字、置中。"""
    set_cell_bg(cell, bg)
    set_cell_border(cell, color=rgb_hex(BLUE_DEEP), size="6")
    cell_write(cell, text, bold=True, size=size, color=WHITE, center=True)


def data_cell(cell, text, row_idx: int = 0, center: bool = False) -> None:
    """資料儲存格：依列號交替底色，提升長表格的可讀性。"""
    set_cell_bg(cell, BLUE_LIGHT if row_idx % 2 == 0 else GRAY_LIGHT)
    set_cell_border(cell, color=rgb_hex(BLUE_MID))
    cell_write(cell, text, center=center)


def section_heading(doc, text, level: int = 1) -> None:
    """帶底線的章節標題。"""
    p = doc.add_paragraph()
    p.clear()
    prefix = "▌" if level == 1 else "▸"
    r = p.add_run(f"{prefix} {text}")
    r.bold = True
    r.font.size = Pt(14 if level == 1 else 12)
    r.font.name = DEFAULT_CJK_FONT
    r.font.color.rgb = BLUE_DEEP
    set_east_asia_font(r)

    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:color"), rgb_hex(BLUE_MID))
    pBdr.append(bot)
    p._p.get_or_add_pPr().append(pBdr)

    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def cover_page(doc, title, subtitle, info_pairs,
               accent_color: RGBColor = BLUE_DEEP) -> None:
    """通用封面頁：主標題 + 副標題 + 兩欄成對的資訊表。"""
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.name = DEFAULT_CJK_FONT
    r.font.color.rgb = accent_color
    set_east_asia_font(r)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(16)
        r2.font.name = DEFAULT_CJK_FONT
        r2.font.color.rgb = BLUE_MID
        set_east_asia_font(r2)

    doc.add_paragraph()

    if info_pairs:
        flat = list(info_pairs.items())
        tbl = doc.add_table(rows=(len(flat) + 1) // 2, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        for col_idx, w in enumerate((Cm(3), Cm(5.5), Cm(3), Cm(5.5))):
            tbl.columns[col_idx].width = w

        for i in range(0, len(flat), 2):
            row = tbl.rows[i // 2]
            k1, v1 = flat[i]
            header_cell(row.cells[0], k1)
            data_cell(row.cells[1], v1, row_idx=i // 2)
            if i + 1 < len(flat):
                k2, v2 = flat[i + 1]
                header_cell(row.cells[2], k2)
                data_cell(row.cells[3], v2, row_idx=i // 2)
            else:
                data_cell(row.cells[2], "", row_idx=i // 2)
                data_cell(row.cells[3], "", row_idx=i // 2)

    doc.add_paragraph()
    doc.add_page_break()


def new_doc_a4(default_font: str = DEFAULT_CJK_FONT):
    """建立標準 A4 文件（邊界：上下 2cm、左右 2.5cm）。"""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = default_font
    style.font.size = Pt(11)
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), default_font)
    except Exception:
        # 少數 python-docx 版本的 Normal style 沒有 rPr，非致命
        pass
    return doc


def add_header_footer(doc, header_text, show_page: bool = True) -> None:
    """加入頁首文字與頁碼欄位。"""
    hp = doc.sections[0].header.paragraphs[0]
    hp.text = header_text
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if hp.runs:
        hp.runs[0].font.size = Pt(9)
        hp.runs[0].font.color.rgb = MUTED_TEXT

    if not show_page:
        return

    fp = doc.sections[0].footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    for tag, text in (("begin", None), (None, " PAGE "), ("end", None)):
        if tag:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), tag)
            run._r.append(fc)
        else:
            it = OxmlElement("w:instrText")
            it.text = text
            run._r.append(it)
